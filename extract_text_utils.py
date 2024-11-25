from utils.config import log_entry_exit
from pdf2image import convert_from_path, convert_from_bytes
from PyPDF2 import PdfReader
import pytesseract
import os
import docx # dox2txt
from PIL import Image
import io, json
from tempfile import SpooledTemporaryFile
from utils.prompt import extract_skills_from_docs_prompt
from integration.openai_integration import get_openai_response
import requests


def extract_text(file_data, file_type): 
    '''
        file_data: SpooledTemporaryFile.read()
        file_type: str
    '''
    if file_type == 'pdf':
        extracted_text = extract_text_from_pdf(file_data)
    elif file_type == 'docx' or file_type == 'doc':
        extracted_text = extract_text_from_docx(file_data)
    elif file_type == 'txt':
        extracted_text = extract_text_from_txt(file_data)
    elif file_type in ['jpg', 'jpeg', 'png']:
        extracted_text = extract_text_from_image(file_data.read())
    elif file_type in ['wav', 'mp3']:
        extracted_text = extract_text_from_audio(file_data)
    else:
        extracted_text = ""

    return extracted_text


def extract_text_from_pdf(pdf_file: SpooledTemporaryFile):
    text = ''
    # extract text from pdf
    try:
        pdf = PdfReader(pdf_file)
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        print(e)

    text += "\n\n"

    # extract text using ocr as well
    # pdf_file.seek(0)
    try:
        pdf = convert_from_bytes(pdf_file)
        for page in pdf:
            text += pytesseract.image_to_string(page) + "\n"
    except Exception as e:
        print(e)
    
    return text


def extract_text_from_image(image_file):
    img = Image.open(io.BytesIO(image_file))
    return pytesseract.image_to_string(img)


def extract_text_from_docx(docx_file):
    text = ''
    doc = docx.Document(io.BytesIO(docx_file))

    for para in doc.paragraphs:
        text += para.text + "\n"

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"

    
    # for images:-
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image = rel.target_part.blob
            img = Image.open(io.BytesIO(image))
            text += pytesseract.image_to_string(img) + "\n"
    
    return text


def extract_text_from_txt(txt_file):
    return io.BytesIO(txt_file).read().decode("utf-8")


def extract_text_from_audio(audio_file: SpooledTemporaryFile):
    try:
        headers = {
            "Ocp-Apim-Subscription-Key": os.environ.get("SPEECH_API_KEY"),
            "Accept": "application/json",
        }

        response = requests.post(
            f"{os.environ.get('SPEECH_TO_TEXT_API_URL')}",
            headers=headers,
            data=audio_file,
        )
        result = response.json()

        if "RecognitionStatus" in result and result["RecognitionStatus"] == "Success":
            return result["DisplayText"]
        else:
            return ""
    except Exception as e:
        print(e)
        return ""



def get_doc_details_llm_call(uploaded_doc):
    try:
        prompt = extract_skills_from_docs_prompt.format(uploaded_doc=uploaded_doc)
        extract_skills = get_openai_response("You are a helpful assistant who generates JSON according to prompt",prompt)

        return json.loads(extract_skills[extract_skills.find("{"):extract_skills.rfind("}")+1])
    except Exception as e:
        raise e
